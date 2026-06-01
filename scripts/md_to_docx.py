#!/usr/bin/env python3
"""Markdown -> DOCX Converter fuer die Projektdokumentation.

Parst die Projektdoku-Markdown (Headings, Listen, Tabellen, Code-Bloecke,
Bold/Inline-Code) und baut ein professionelles Word-Dokument mit:
- Styled Heading-Hierarchie (H1-H4)
- Auto-aktualisierendem Word-TOC-Feld (statt der manuellen MD-Inhaltsverzeichnis-Liste)
- Monospace + grau hinterlegten Code-Bloecken
- Gerahmten Tabellen mit Header-Shading

Aufruf: python scripts/md_to_docx.py <input.md> <output.docx> [titel-override]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x2A, 0x44)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
CODE_BG = "F0F2F5"
HDR_BG = "1F2A44"


def _set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _add_toc(doc):
    """Fuegt ein Word-TOC-Feld ein (aktualisiert sich via F9 / beim Oeffnen)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "Inhaltsverzeichnis — in Word mit F9 aktualisieren."
    run_ph = OxmlElement("w:r"); run_ph.append(placeholder)
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep)
    p._p.append(run_ph); r2 = OxmlElement("w:r"); r2.append(fld_end); p._p.append(r2)


_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_inline(paragraph, text, base_mono=False):
    """Fuegt Text mit **bold** + `code`-Inline-Formatierung als Runs hinzu."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
        else:
            run = paragraph.add_run(part)
        if base_mono:
            run.font.name = "Consolas"; run.font.size = Pt(9.5)


def _strip_md_links(text):
    """[Label](#anchor) -> Label."""
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)


def convert(md_path, docx_path, title_override=None):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Basis-Style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []
    skip_toc_section = False
    toc_done = False

    while i < n:
        line = lines[i]

        # Code-Fence
        if line.strip().startswith("```"):
            if in_code:
                # Code-Block schreiben
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), CODE_BG)
                pPr.append(shd)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"; run.font.size = Pt(9)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # TOC-Sektion der MD ueberspringen (durch Word-TOC ersetzt)
        if re.match(r"^##\s+Inhaltsverzeichnis", line):
            skip_toc_section = True; i += 1; continue
        if skip_toc_section:
            if re.match(r"^---\s*$", line) or re.match(r"^##\s+\d", line) or re.match(r"^##\s+[1-9]", line):
                skip_toc_section = False
                # nicht 'continue' — diese Zeile normal verarbeiten (falls Heading)
            else:
                i += 1; continue

        # Horizontale Linie
        if re.match(r"^---\s*$", line):
            i += 1; continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1)); text = _strip_md_links(m.group(2).strip())
            if level == 1 and title_override:
                text = title_override
            h = doc.add_heading(level=level)
            _add_inline(h, text)
            for run in h.runs:
                run.font.color.rgb = NAVY if level <= 2 else ACCENT
            # Nach dem Titel (erstes H1): Word-TOC-Feld + Seitenumbruch
            if level == 1 and not toc_done:
                toc_done = True
                doc.add_paragraph()
                _add_toc(doc)
                doc.add_page_break()
            i += 1; continue

        # Tabelle (Zeile mit | + naechste Zeile Separator)
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i+1]) and "-" in lines[i+1]:
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            tbl = doc.add_table(rows=1, cols=len(header))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = "Table Grid"
            for ci, htext in enumerate(header):
                cell = tbl.rows[0].cells[ci]
                cell.text = ""
                _add_inline(cell.paragraphs[0], _strip_md_links(htext))
                for run in cell.paragraphs[0].runs:
                    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, HDR_BG)
            for r in rows:
                cells = tbl.add_row().cells
                for ci in range(len(header)):
                    val = r[ci] if ci < len(r) else ""
                    cells[ci].text = ""
                    _add_inline(cells[ci].paragraphs[0], _strip_md_links(val))
            doc.add_paragraph()
            i = j; continue

        # Bullet-Liste (- oder *), inkl. Sub-Indent
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)); lvl = min(indent // 2, 2)
            style = "List Bullet" if lvl == 0 else f"List Bullet {lvl+1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, _strip_md_links(m.group(2).strip()))
            i += 1; continue

        # Nummerierte Liste
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            _add_inline(p, _strip_md_links(m.group(2).strip()))
            i += 1; continue

        # Leerzeile
        if not line.strip():
            i += 1; continue

        # Normaler Absatz
        p = doc.add_paragraph()
        _add_inline(p, _strip_md_links(line.strip()))
        i += 1

    # TOC nach dem ersten Heading (Titel) einfuegen
    # -> wir fuegen es am Anfang nach Titel ein: finde erstes Heading-Paragraph
    # Einfacher: TOC-Feld als zweites Element. Wir bauen Doc neu-geordnet:
    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else None
    out = convert(src, dst, title)
    print(f"OK: {out}")
